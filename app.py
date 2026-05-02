import io
import json
import os
import re
from pathlib import Path
from typing import Any, Dict, Tuple

from anthropic import Anthropic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from dotenv import load_dotenv
from flask import Flask, jsonify, render_template, request, send_file
from fpdf import FPDF


SYSTEM_PROMPT = """
You are an expert resume writer with 15 years of experience crafting 
ATS-optimized, visually polished resumes for technical roles. You have 
deep expertise in helping candidates pass automated ATS bots, AI 
screening tools, and human HR reviewers to land interviews.

Your goal is to tailor the candidate's resume to the provided job 
description while matching a strict visual format, maximizing ATS 
score, and fitting within a strict 1-page length constraint.

═══════════════════════════════════════
CANDIDATE WORK AUTHORIZATION STATUS
═══════════════════════════════════════
The candidate is on an F1 Student Visa (OPT/STEM OPT).
This means:
  ✅ Can work in the US without employer visa sponsorship
  ✅ Eligible for roles stating "must be authorized to work in the US"
  ✅ Eligible for roles with no sponsorship/citizenship mention at all
  ❌ NOT a US Citizen
  ❌ NOT a Green Card / Permanent Resident holder
  ❌ NOT eligible for any role requiring security clearance of any level
     (Secret, Top Secret, TS/SCI, DoD, DoE, DHS — ALL are hard blocks)

Use this status to drive the eligibility logic in STEP 0 below.

═══════════════════════════════════════
INPUTS
═══════════════════════════════════════
CANDIDATE RESUME:
{resume}

TARGET JOB DESCRIPTION:
{job_description}

═══════════════════════════════════════
STEP 0 — ELIGIBILITY CHECK (RUN FIRST)
═══════════════════════════════════════
Before doing ANYTHING else, scan the FULL job description carefully
and classify it into one of three buckets.

─────────────────────────────────────
BUCKET A — HARD BLOCK ❌
─────────────────────────────────────
Trigger if the JD contains ANY of these (exact or semantic match):
  - "US Citizenship required" / "Must be a US Citizen"
  - "Green Card required" / "Permanent Resident required"
  - "Only US Citizens and Green Card holders eligible"
  - "Secret Clearance required"
  - "Top Secret / TS / TS-SCI clearance required"
  - "Active or ability to obtain Security Clearance"
  - "Must hold or be eligible for government/DoD/DoE/DHS clearance"
  - "Must be clearable"

ACTION: STOP immediately. Do NOT generate a resume under any circumstance.
Return ONLY this output:

---ELIGIBILITY CHECK FAILED---

⚠️ This job is NOT suitable for your profile.

Reason: This position requires [exact phrase found in JD], which is 
restricted to US Citizens or Permanent Residents only.

Your Status: F1 Visa (OPT/STEM OPT) — you do not qualify for this 
role regardless of sponsorship or work authorization.

Recommendation: Skip this application. Focus on roles that state:
  ✅ "Open to all work authorizations"
  ✅ "Will sponsor H-1B"
  ✅ "F1/OPT candidates welcome"
  ✅ No citizenship, Green Card, or clearance restrictions mentioned

─────────────────────────────────────
BUCKET B — ELIGIBLE ✅
─────────────────────────────────────
Trigger if the JD contains ANY of these OR no restrictions at all:
  - "Must be authorized to work in the US"
  - "No visa sponsorship available"
  - "Must be able to work without sponsorship"
  - "Employment authorization required"
  - No mention of citizenship, Green Card, or clearance anywhere

ACTION: Candidate qualifies — F1 OPT/STEM OPT satisfies these 
requirements. Proceed silently to STEP 1.
Do NOT mention the eligibility check anywhere in the resume output.

─────────────────────────────────────
BUCKET C — AMBIGUOUS ⚠️
─────────────────────────────────────
Trigger if the JD has unclear or mixed signals such as:
  - "Preferred: US Citizen or Green Card" (preferred, NOT required)
  - "Public Trust clearance required" (sometimes allows non-citizens)
  - "Security clearance a plus" / "clearance preferred"
  - Sponsorship language is vague or contradictory

ACTION: Do NOT block. Generate the full resume AND prepend this note:

---ELIGIBILITY NOTE---

⚠️ Advisory: This job contains unclear work authorization language:
"[exact phrase from JD]"

Your Status: F1 Visa (OPT/STEM OPT) — you CAN work without 
sponsorship but are NOT a US Citizen or Green Card holder.

Recommendation: Apply but verify directly with the recruiter whether 
F1 OPT candidates are considered before investing time in interviews.

---RESUME---
[Resume follows below]

═══════════════════════════════════════
STEP 1 — PRE-RESUME ANALYSIS (SILENT)
═══════════════════════════════════════
Before writing anything, silently perform this analysis:

1. TARGET ROLE IDENTIFICATION:
   - Identify the single primary role this JD is hiring for
   - Extract the top 3 domain areas emphasized in the JD
   - Note the seniority level (entry, mid, senior)

2. KEYWORD EXTRACTION:
   - Extract ALL technical keywords from the JD
   - Separate into: tools, languages, methodologies, soft skills
   - Flag which keywords already exist in the candidate's resume
   - Flag which keywords are missing but can be naturally added

3. RELEVANCE MAPPING:
   - Map each candidate experience bullet to JD requirements
   - Identify which projects are most relevant (max keep 2–3)
   - Identify which projects to prune (irrelevant to this JD)
   - Flag any weak bullets that need stronger action verbs

4. GAP IDENTIFICATION:
   - Note any JD requirements completely missing from the resume
   - DO NOT fabricate these — flag them silently, do not add them

This analysis drives all decisions in STEP 2 below.
Do NOT output this analysis — it is for internal reasoning only.

═══════════════════════════════════════
STEP 2 — PROFILE SUMMARY (CRITICAL)
═══════════════════════════════════════
Write a FRESH, UNIQUE 4-sentence profile summary for EVERY job 
description. NEVER reuse or recycle a previous profile summary.
The summary must be re-written from scratch to mirror this specific 
JD's language, priorities, and target role.

Follow this EXACT sentence-by-sentence structure:

SENTENCE 1 — WHO YOU ARE:
  Formula: [Degree] + [Years of Experience] + [Top 2–3 domain areas 
            from THIS specific JD]
  Purpose: Establishes identity and seniority. Names the exact role.
  Rule:    Extract domain areas verbatim from the JD — no synonyms.
           Lead with the exact job title or closest match.
  Example: "Master's in Data Science graduate with **2 years of 
            experience** in **data analysis**, **machine learning**, 
            and **business intelligence**."

SENTENCE 2 — WHAT YOU KNOW (Tools & Tech):
  Formula: [Proficiency statement] + [Tools from JD that exist in 
            candidate resume] + [Platforms/Frameworks from JD]
  Purpose: ATS sentence — maximum keyword density from the JD.
  Rule:    ONLY include tools present in BOTH the resume AND the JD.
           Prioritize JD tools. Never include tools not in the resume.
  Example: "Proficient in **Python**, **SQL**, **Tableau**, and 
            **Power BI**, with hands-on exposure to **Big Data** 
            technologies including **Hadoop**, **Spark**, and **Hive**."

SENTENCE 3 — WHAT YOU'VE DONE (Skills in Action):
  Formula: [Experience areas] + [Competencies] + [Technical focus]
  Purpose: Bridges tools to real-world application. Shows depth.
  Rule:    Use strong action-oriented phrases.
           Reflect the key RESPONSIBILITIES listed in the JD.
           Do NOT repeat tools already mentioned in Sentence 2.
           Include at least one quantified achievement if available.
  Example: "Experienced in **predictive modelling**, **statistical 
            analysis**, **ETL pipeline development**, and 
            **cloud-based analytics** with a focus on **anomaly 
            detection** and **data visualization**."

SENTENCE 4 — WHY YOU DO IT (Passion + Value):
  Formula: [Mission statement] + [Company or role impact alignment]
  Purpose: Humanizes the resume. Signals cultural fit.
  Rule:    Research the JD for company mission or role impact 
           statements and mirror them. Avoid generic filler phrases 
           like "hardworking", "passionate learner", or "team player".
  Example: "Passionate about leveraging **AI/ML** to derive 
            actionable business insights and solve real-world 
            challenges at scale."

PROFILE SUMMARY STYLE RULES:
  - Dense paragraph format — NO bullet points in this section
  - NO first-person pronouns ("I", "my", "me") — implied third person
  - Bold ALL technical terms, tools, domain phrases, and numbers
  - 3–5 lines MAX — recruiters scan in 6 seconds, not read
  - Every word must mirror this JD's exact language
  - ZERO fabrication — only use what exists in the candidate's resume
  - NEVER copy the example sentences above — write fresh every time

═══════════════════════════════════════
STEP 3 — CONTENT QUALITY RULES
═══════════════════════════════════════

1. BULLET POINT FORMULA (MANDATORY FOR ALL BULLETS):
   Every bullet MUST follow this structure:
   [Strong Action Verb] + [What You Did] + [Tool/Method Used] 
   + [Quantified Result or Business Impact]

   Strong action verbs to use:
   Architected, Engineered, Optimized, Automated, Deployed,
   Developed, Designed, Implemented, Spearheaded, Delivered,
   Built, Reduced, Improved, Increased, Streamlined

   Weak verbs to NEVER use:
   Worked, Helped, Assisted, Leveraged, Utilized, Did, Made

   ❌ WEAK:  "Worked in Agile environments with product managers"
   ✅ STRONG: "Collaborated across 3-team Agile sprints to deploy 
               2 production pipelines, reducing release time by 20%"

2. METRICS — PRESERVE AND ENHANCE:
   Always preserve these exact metrics from the resume:
   - 4.0 GPA
   - 30% query execution time reduction
   - 95% report accuracy improvement
   - 20+ clients
   If a bullet has no metric, add business impact context instead.

3. TECHNICAL SKILLS SECTION RULES:
   - Group by category (already done in reference)
   - Remove "familiar" or "basic" qualifiers — own it or drop it
   - Add specificity: "Python" → "Python (Pandas, NumPy, Scikit-learn)"
   - Only include skills present in the candidate's resume

4. PROJECTS — STRICT SELECTION:
   - Maximum 2–3 projects, chosen based on JD relevance
   - Prune any project with zero relevance to the target role
   - Each project needs minimum 2 strong bullets with impact
   - Project titles: Bold + Title Case (NOT all caps)
     ✅ **Retail Sales Prediction**
     ❌ **RETAIL SALES PREDICTION**

5. CERTIFICATIONS:
   - Only include if relevant to the JD or recent (within 3 years)
   - If certification is older than 3 years and low-value, prune it

6. HOBBIES:
   - Only include if they signal relevant skills or traits
   - Remove generic entries like "playing games" or vague reading
   - Keep: Chess, Open Source Contributions, AI/ML research reading
   - Remove: Anything that doesn't add professional signal

═══════════════════════════════════════
STEP 4 — VISUAL FORMATTING RULES
═══════════════════════════════════════

1. MARGINS & SPACING:
   - Narrow Margin layout (0.5 inch). Dense content.
   - NO empty lines between section content, separator, and 
     the next section header.

2. CONTACT HEADER (CENTERED):
   - Candidate full name: Bold, Centered, larger font — Line 1
   - Line 2 (centered, single line):
     City, State | Email | Phone | LinkedIn URL | GitHub URL
   - Full-width horizontal line immediately after: 
     "_______________________________________________"

3. SECTION HEADERS — EXACT FORMAT (case and colon sensitive):
   Use these EXACT titles, no deviation:
   PROFILE                    ← ALL CAPS, Bold, Underlined, no colon
   EXPERIENCE:                ← ALL CAPS, Bold, Underlined, colon attached
   EDUCATION                  ← ALL CAPS, Bold, Underlined, no colon
   LEADERSHIP:                ← ALL CAPS, Bold, Underlined, colon attached
   TECHNICAL SKILLS           ← ALL CAPS, Bold, Underlined, no colon
   ACADEMIC PROJECTS          ← ALL CAPS, Bold, Underlined, no colon
   ADDITIONAL CERTIFICATIONS  ← ALL CAPS, Bold, Underlined, no colon
   Hobbies:                   ← Title Case, Bold, Underlined, colon attached

4. HORIZONTAL SEPARATORS:
   - Full-width line after contact header AND after every section
   - Format: "_______________________________________________"
   - NO blank lines before or after separator

5. EXPERIENCE & LEADERSHIP STRUCTURE:
   - Line 1: **Organization – Location** (Bold)
   - Line 2: **Role | MMM YYYY – MMM YYYY** (Bold)
   - Bullets: use "•" filled circle (not "-" dash)
   - Bold all technical terms, tools, and metrics inline

6. DATE FORMAT CONSISTENCY:
   Use abbreviated month format universally across ALL sections:
   ✅ Aug 2022 – Aug 2024
   ❌ August 2022 – August 2024
   ❌ 2022 – 2024 (missing months)
   ❌ Mixing any formats

7. EDUCATION STRUCTURE (SINGLE TAB-ALIGNED LINE):
   **University Name –** City, State    **Degree** | GPA: X.X    YYYY–YYYY
   Example:
   **Pace University –** New York, NY    **Master's in Data Science** | GPA: 4.0    2024–2026

8. ACADEMIC PROJECT TITLE FORMAT:
   Bold + Title Case — NOT ALL CAPS
   ✅ **Retail Sales Prediction**
   ✅ **AI-Powered Cybersecurity Data Analysis Platform**
   ❌ **RETAIL SALES PREDICTION**

═══════════════════════════════════════
STEP 5 — LENGTH & PRUNING RULES
═══════════════════════════════════════
- MAX LENGTH: Strictly 1 page — no exceptions
- PRUNING ORDER (remove in this order until 1 page fits):
    1. Least relevant academic projects
    2. Weakest/most generic experience bullets
    3. Low-value certifications (old or irrelevant)
    4. Generic hobbies with no professional signal
    5. Leadership bullets beyond top 2
- KEEP ALWAYS: Core metrics, GPA, top 2 projects, all experience

═══════════════════════════════════════
STEP 6 — OUTPUT FORMAT
═══════════════════════════════════════
Return response in this EXACT structure with NO deviations:

---RESUME---
[Full tailored 1-page resume following all rules above]

---KEYWORDS MATCHED---
- [Keyword 1] — [where it was used in the resume]
- [Keyword 2] — [where it was used in the resume]
[Continue for all matched keywords]

---MISSING KEYWORDS---
- [Keyword from JD that could NOT be added — not in candidate resume]
[List only — do not fabricate or add these to the resume]

---ATS SCORE---
Score: [X/100]

Breakdown:
- Keyword Match Rate (40 pts): [X/40] — [explanation]
- Formatting & Readability (30 pts): [X/30] — [explanation]
- Experience Relevance (30 pts): [X/30] — [explanation]

Overall Verdict: [2–3 sentence honest assessment of resume 
strength for this specific JD and what would improve the score]
"""


MODEL_NAME = "claude-sonnet-4-6"
ALLOWED_FONT_SIZES = {8, 10, 12}
INPUT_TOKEN_PRICE = 0.000001
OUTPUT_TOKEN_PRICE = 0.000005

BASE_DIR = Path(__file__).resolve().parent
DATA_FILE = BASE_DIR / "data.json"

app = Flask(__name__)
load_dotenv(BASE_DIR / ".env", override=True)


def _default_settings() -> Dict[str, Any]:
    return {"base_resume": "", "font_size": 10}


def normalize_font_size(value: Any) -> int:
    try:
        parsed = int(value)
    except (TypeError, ValueError):
        return 10
    return parsed if parsed in ALLOWED_FONT_SIZES else 10


def load_settings() -> Dict[str, Any]:
    default = _default_settings()
    if not DATA_FILE.exists():
        return default

    try:
        with DATA_FILE.open("r", encoding="utf-8") as file:
            loaded = json.load(file)
    except (json.JSONDecodeError, OSError):
        return default

    base_resume = str(loaded.get("base_resume", ""))
    font_size = normalize_font_size(loaded.get("font_size", 10))
    return {"base_resume": base_resume, "font_size": font_size}


def save_settings(settings: Dict[str, Any]) -> None:
    serializable = {
        "base_resume": str(settings.get("base_resume", "")),
        "font_size": normalize_font_size(settings.get("font_size", 10)),
    }
    with DATA_FILE.open("w", encoding="utf-8") as file:
        json.dump(serializable, file, indent=2)


def parse_model_output(text: str) -> Dict[str, str]:
    result = {"resume": "", "keywords": "", "ats_score": ""}
    cleaned = text.strip()

    resume_header = "---RESUME---"
    keywords_header = "---KEYWORDS MATCHED---"
    missing_header = "---MISSING KEYWORDS---"
    ats_header = "---ATS SCORE---"

    keywords_pos = cleaned.find(keywords_header)
    missing_pos = cleaned.find(missing_header)
    ats_pos = cleaned.find(ats_header)

    resume_end_candidates = [pos for pos in (keywords_pos, missing_pos, ats_pos) if pos != -1]
    resume_search_end = min(resume_end_candidates) if resume_end_candidates else len(cleaned)
    resume_pos = cleaned.rfind(resume_header, 0, resume_search_end)
    if resume_pos != -1:
        resume_start = resume_pos + len(resume_header)
        resume_end = min([pos for pos in resume_end_candidates if pos > resume_start], default=len(cleaned))
        result["resume"] = cleaned[resume_start:resume_end].strip()

    if keywords_pos != -1:
        keywords_start = keywords_pos + len(keywords_header)
        keywords_end = min(
            [pos for pos in (missing_pos, ats_pos) if pos != -1 and pos > keywords_start],
            default=len(cleaned),
        )
        result["keywords"] = cleaned[keywords_start:keywords_end].strip()

    if ats_pos != -1:
        ats_start = ats_pos + len(ats_header)
        result["ats_score"] = cleaned[ats_start:].strip()

    if not result["resume"]:
        result["resume"] = cleaned
    result["resume"] = normalize_resume_markup(result["resume"])

    return result


def compute_cost(usage: Any) -> Dict[str, Any]:
    input_tokens = 0
    output_tokens = 0

    if usage is not None:
        input_tokens = int(getattr(usage, "input_tokens", 0) or 0)
        output_tokens = int(getattr(usage, "output_tokens", 0) or 0)

    input_cost = input_tokens * INPUT_TOKEN_PRICE
    output_cost = output_tokens * OUTPUT_TOKEN_PRICE
    total_cost = input_cost + output_cost

    return {
        "input_tokens": input_tokens,
        "output_tokens": output_tokens,
        "input_cost": round(input_cost, 8),
        "output_cost": round(output_cost, 8),
        "total_cost": round(total_cost, 8),
        "display": f"${total_cost:.8f}",
    }


def get_anthropic_client() -> Anthropic:
    api_key = os.environ.get("ANTHROPIC_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("Missing ANTHROPIC_API_KEY environment variable.")
    return Anthropic(api_key=api_key)


def build_messages(base_resume: str, job_description: str) -> list[Dict[str, Any]]:
    return [
        {
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": f"BASE_RESUME:\n{base_resume}",
                    "cache_control": {"type": "ephemeral"},
                },
                {
                    "type": "text",
                    "text": f"JOB_DESCRIPTION:\n{job_description}",
                },
            ],
        }
    ]


def sanitize_pdf_text(text: str) -> str:
    replacements = {
        "\u2022": "-",
        "\u2013": "-",
        "\u2014": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\xa0": " ",
    }
    normalized = text
    for source, target in replacements.items():
        normalized = normalized.replace(source, target)
    return normalized.encode("latin-1", errors="replace").decode("latin-1")


def split_markdown_bold_segments(text: str) -> list[tuple[str, bool]]:
    segments: list[tuple[str, bool]] = []
    parts = text.split("**")
    for idx, part in enumerate(parts):
        if not part:
            continue
        segments.append((part, idx % 2 == 1))
    if not segments:
        return [("", False)]
    return segments


def normalize_resume_markup(text: str) -> str:
    normalized = text
    normalized = normalized.replace("&nbsp;", " ")
    normalized = normalized.replace("&amp;nbsp;", " ")
    normalized = normalized.replace("&#160;", " ")
    normalized = normalized.replace("&#xA0;", " ")
    normalized = normalized.replace("&#xa0;", " ")
    normalized = re.sub(
        r"(?is)(?:<u>|&lt;u&gt;)\s*(.*?)\s*(?:</u>|&lt;/u&gt;)",
        r"**\1**",
        normalized,
    )
    normalized = re.sub(
        r"(?is)(?:<ins>|&lt;ins&gt;)\s*(.*?)\s*(?:</ins>|&lt;/ins&gt;)",
        r"**\1**",
        normalized,
    )
    normalized = re.sub(r"(?i)</?u>", "", normalized)
    normalized = re.sub(r"(?i)&lt;/?u&gt;", "", normalized)
    normalized = re.sub(r"(?i)</?ins>", "", normalized)
    normalized = re.sub(r"(?i)&lt;/?ins&gt;", "", normalized)
    normalized = re.sub(r"\*{4,}\s*(.*?)\s*\*{4,}", r"**\1**", normalized)
    return normalized


def generate_pdf_bytes(resume_text: str, font_size: int) -> bytes:
    resume_text = normalize_resume_markup(resume_text)
    pdf = FPDF()
    narrow_margin_mm = 12.7
    pdf.set_margins(narrow_margin_mm, narrow_margin_mm, narrow_margin_mm)
    pdf.set_auto_page_break(auto=True, margin=narrow_margin_mm)
    pdf.add_page()
    pdf.set_font("Times", size=font_size)

    line_height = max(3.2, font_size * 0.42)
    rendered_line_count = 0
    for line in resume_text.splitlines():
        stripped = sanitize_pdf_text(line.strip())
        if not stripped:
            continue
        plain_for_center = stripped.replace("**", "")
        if rendered_line_count < 2:
            pdf.set_font("Times", style="B", size=font_size)
            pdf.cell(0, 10, plain_for_center, align="C", ln=1)
            pdf.set_font("Times", size=font_size)
            rendered_line_count += 1
            continue
        if stripped.replace("_", "") == "" and len(stripped) >= 10:
            y_pos = min(pdf.get_y() + 2, pdf.h - pdf.b_margin)
            pdf.set_y(y_pos)
            pdf.line(pdf.l_margin, y_pos, pdf.w - pdf.r_margin, y_pos)
            pdf.set_y(y_pos + 2)
            pdf.set_x(pdf.l_margin)
            rendered_line_count += 1
            continue
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, line_height, stripped, markdown=True)
        rendered_line_count += 1

    payload = pdf.output(dest="S")
    if isinstance(payload, (bytes, bytearray)):
        return bytes(payload)
    return payload.encode("latin-1", errors="replace")


def generate_docx_bytes(resume_text: str, font_size: int) -> bytes:
    resume_text = normalize_resume_markup(resume_text)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(font_size)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    rendered_line_count = 0
    for line in resume_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.replace("_", "") == "" and len(stripped) >= 10:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0

            p_pr = paragraph._p.get_or_add_pPr()
            p_bdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "0")
            bottom.set(qn("w:color"), "auto")
            p_bdr.append(bottom)
            p_pr.append(p_bdr)

            rendered_line_count += 1
            continue
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        if rendered_line_count < 2:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for segment_text, segment_bold in split_markdown_bold_segments(stripped):
            run = paragraph.add_run(segment_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(font_size)
            run.bold = segment_bold or (rendered_line_count < 2)
        rendered_line_count += 1

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/settings", methods=["GET"])
def get_settings():
    return jsonify(load_settings())


@app.route("/settings", methods=["POST"])
def update_settings():
    data = request.get_json(silent=True) or {}
    base_resume = str(data.get("base_resume", "")).strip()
    font_size = normalize_font_size(data.get("font_size", 10))

    payload = {"base_resume": base_resume, "font_size": font_size}
    save_settings(payload)
    return jsonify(payload)


@app.route("/tailor", methods=["POST"])
def tailor_resume():
    data = request.get_json(silent=True) or {}
    base_resume = str(data.get("base_resume", "")).strip()
    job_description = str(data.get("job_description", "")).strip()

    if not base_resume:
        return jsonify({"error": "Base resume cannot be empty."}), 400
    if not job_description:
        return jsonify({"error": "Job description cannot be empty."}), 400

    try:
        client = get_anthropic_client()
        response = client.messages.create(
            model=MODEL_NAME,
            max_tokens=1800,
            system=SYSTEM_PROMPT,
            messages=build_messages(base_resume, job_description),
        )
        response_text = "".join(
            block.text for block in getattr(response, "content", []) if block.type == "text"
        ).strip()
        parsed = parse_model_output(response_text)
        costs = compute_cost(getattr(response, "usage_metadata", None))
        if costs["input_tokens"] == 0 and costs["output_tokens"] == 0:
            costs = compute_cost(getattr(response, "usage", None))

        return jsonify(
            {
                "resume": parsed["resume"],
                "keywords_matched": parsed["keywords"],
                "ats_score": parsed["ats_score"],
                "raw_response": response_text,
                "cost": costs,
            }
        )
    except Exception as exc:
        error_message = str(exc).strip() or "Unknown API error."
        if "timeout" in error_message.lower():
            return (
                jsonify({"error": "Anthropic API timeout. Please try again."}),
                500,
            )
        return jsonify({"error": f"Failed to tailor resume: {error_message}"}), 500


def _validate_download_input(data: Dict[str, Any]) -> Tuple[str, int, Any]:
    resume_text = str(data.get("resume_text", "")).strip()
    font_size = normalize_font_size(data.get("font_size", 10))
    if not resume_text:
        return "", font_size, (jsonify({"error": "Resume text is required."}), 400)
    return resume_text, font_size, None


@app.route("/download/pdf", methods=["POST"])
def download_pdf():
    data = request.get_json(silent=True) or {}
    resume_text, font_size, error_response = _validate_download_input(data)
    if error_response:
        return error_response

    try:
        pdf_bytes = generate_pdf_bytes(resume_text, font_size)
        return send_file(
            io.BytesIO(pdf_bytes),
            as_attachment=True,
            download_name="tailored_resume.pdf",
            mimetype="application/pdf",
        )
    except Exception as exc:
        return jsonify({"error": f"Failed to generate PDF: {str(exc).strip()}"}), 500


@app.route("/download/docx", methods=["POST"])
def download_docx():
    data = request.get_json(silent=True) or {}
    resume_text, font_size, error_response = _validate_download_input(data)
    if error_response:
        return error_response

    docx_bytes = generate_docx_bytes(resume_text, font_size)
    return send_file(
        io.BytesIO(docx_bytes),
        as_attachment=True,
        download_name="tailored_resume.docx",
        mimetype=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )


if __name__ == "__main__":
    if not DATA_FILE.exists():
        save_settings(_default_settings())
    app.run(debug=True)
